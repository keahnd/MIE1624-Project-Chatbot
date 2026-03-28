import json
from pathlib import Path
from datetime import datetime
from docx import Document

from basic_crewai_1.crew import BasicCrewai1


REPORT_PATH = "data/final_report.docx"
RAI_CONTEXT_PATH = r"C:\Users\keahn\Documents\UofT\Coursework\MIE1624\Project\rai_analysis\outputs\rai_context.json"
OUTPUT_DIR = "outputs"


def load_rai_context() -> str:
    """Load structured RAI analysis data from rai_context.json."""
    context_path = Path(RAI_CONTEXT_PATH)
    if not context_path.exists():
        return ""
    with open(context_path, encoding="utf-8") as f:
        return json.dumps(json.load(f), indent=2)


def read_docx(path: str) -> str:
    """Read text from a .docx file, including paragraphs and tables."""
    file_path = Path(path)
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    doc = Document(file_path)

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    table_lines = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                table_lines.append(" | ".join(cells))

    parts = []
    if paragraphs:
        parts.append("\n".join(paragraphs))
    if table_lines:
        parts.append("\n".join(table_lines))

    full_text = "\n\n".join(parts).strip()

    if not full_text:
        raise ValueError(f"No readable text found in: {file_path}")

    return full_text


def save_answer(question_number: int, question: str, answer: str) -> Path:
    """Save each CrewAI response to a markdown file automatically."""
    output_dir = Path(OUTPUT_DIR)
    output_dir.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    file_path = output_dir / f"qa_{question_number}_{timestamp}.md"

    content = (
        f"# Question {question_number}\n\n"
        f"## User Question\n{question}\n\n"
        f"## CrewAI Answer\n{answer}\n"
    )

    file_path.write_text(content, encoding="utf-8")
    return file_path


def run():
    report_text = read_docx(REPORT_PATH)
    rai_context = load_rai_context()

    if rai_context:
        full_context = f"=== GROUP REPORT ===\n{report_text}\n\n=== RAI ANALYSIS DATA ===\n{rai_context}"
    else:
        full_context = report_text

    print("\nCrewAI chatbot is ready. Type 'stop' to exit.\n")

    i = 1
    while True:
        user_question = input(f"Question {i}: ").strip()

        if not user_question:
            print("Please enter a question.\n")
            continue

        if user_question.lower() in ("stop", "exit", "quit"):
            print("Session ended.")
            break

        inputs = {
            "full_context": full_context,
            "user_question": user_question
        }

        result = BasicCrewai1().crew().kickoff(inputs=inputs)
        result_text = str(result)

        saved_file = save_answer(i, user_question, result_text)

        print("\n===== ANSWER =====\n")
        print(result_text)
        print(f"\nSaved to: {saved_file}\n")

        i += 1

    print("Session ended.")


def train():
    print("Training hook not implemented.")


def replay():
    print("Replay hook not implemented.")


def test():
    print("Test hook not implemented.")


def run_with_trigger():
    run()